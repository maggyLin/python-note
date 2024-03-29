from docx import Document

# document = Document("D:/project/parseTestplan/5940 CFD_Production Test Plan_C0464_20231011_test.docx")
document = Document("../5940 CFD_Production Test Plan_C0464_20231011_test.docx")

tables = document.tables
table = tables[1] # 抓取第幾個table

# cell(row,column)
for i in range(0,len(table.rows)):
    result = table.cell(i,0).text
    print( result )
    result1 = table.cell(i,1).text
    print( result1)
    result2 = table.cell(i,2).text
    print( result2)
    # result3 = table.cell(i,3).text
    # print( result3)
    # result4 = table.cell(i,4).text
    # print( result4)
    

# cell = len(table.columns)
# print(cell)